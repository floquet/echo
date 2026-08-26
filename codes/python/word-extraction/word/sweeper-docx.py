#!/opt/local/bin/python3.13
# =============================================================================
# sweeper-docx.py
#
# Purpose
# -------
#
# Explore the internal structure of Microsoft Word (*.docx) files.
#
# This program is intended as a reverse-engineering tool rather than a final
# harvesting application.
#
# The goal is to determine:
#
#     • where visible text is stored,
#     • whether user-entered information resides in paragraphs,
#       tables, text boxes, or content controls,
#     • what metadata accompanies each field,
#     • and how future harvesting routines should extract that data.
#
# Lessons Learned
# ---------------
#
# 1. A *.docx file is a ZIP archive containing XML.
#
# 2. python-docx provides a convenient high-level object model but does not
#    expose every XML construct.
#
# 3. Government assessment forms frequently store user input inside
#    Structured Document Tags (<w:sdt>) rather than ordinary paragraphs.
#
# 4. Reverse engineering should begin with an XML sweep before writing a
#    specialized harvesting routine.
# =============================================================================

# =============================================================================
# High-Level python-docx View
#
# python-docx presents a simplified object model consisting primarily of
# paragraphs, tables, sections, styles, and runs.
#
# This abstraction is convenient for editing documents, but it intentionally
# hides much of Word's internal XML representation.
#
# Government forms frequently store user-entered data in structures that are
# NOT exposed through doc.paragraphs.
# =============================================================================

# =============================================================================
# WordprocessingML Namespace
#
# A *.docx file is a ZIP archive containing a collection of XML documents.
# Nearly all document content is stored using Microsoft's WordprocessingML
# vocabulary. XML element names such as <w:p>, <w:t>, and <w:sdt> belong to
# the namespace identified below.
#
# XPath expressions use the symbolic prefix "w". The namespace dictionary
# maps that prefix to the full URI required by the XML parser.
# =============================================================================

# Daniel Topa
# BAE 899149

import os
import sys
import platform
import datetime
from pathlib import Path
from docx    import Document

Word_NameSpace = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
}

# =============================================================================
# paragraph_text()
#
# Collect all visible text beneath an XML element.
#
# Word stores text inside <w:t> ("text") elements. A paragraph, content
# control, or table cell may contain many nested XML objects before reaching
# the actual text nodes.
#
# Rather than depending on XPath, we iterate over every descendant <w:t>
# element and concatenate the text in document order.
#
# This routine works equally well for ordinary paragraphs, content controls,
# table cells, and many other Word objects.
# =============================================================================

def paragraph_text( element ) -> str:
    """Return all Word text contained beneath an XML element."""

    text_tag = f"{{{Word_NameSpace['w']}}}t"

    return "".join(
        node.text or ""
        for node in element.iter(text_tag)
    )

def sweep_document ( filename: str ) -> None:
    path = Path ( filename )

    if not path.exists():
        raise FileNotFoundError( path )

    doc = Document( path )

    print ( "\n=== HIGH-LEVEL PYTHON-DOCX OBJECTS ===" )
    print ( f"Paragraphs       : {len( doc.paragraphs )}" )
    print ( f"Tables           : {len( doc.tables )}" )
    print ( f"Sections         : {len( doc.sections )}" )

# =============================================================================
# XML Paragraph Sweep
#
# Walk every XML paragraph (<w:p>) in the document.
#
# This is a lower-level view than doc.paragraphs. It reveals paragraphs that
# python-docx does not necessarily expose as high-level Paragraph objects.
#
# Many reverse-engineering tasks begin by sweeping every paragraph and
# examining the recovered text.
# =============================================================================

    print ( "\n=== ALL XML PARAGRAPHS ===" )

    #for index, element in enumerate(doc.element.xpath(".//w:p" )):
    for index, element in enumerate( doc.element.xpath( ".//w:p" ) ):
        #print ( f"Element type : ", type(element) ) 
        text = paragraph_text( element )

        if text.strip():
            print ( f"{index:4d}: {text!r}" )

# =============================================================================
# Text Boxes
#
# Many Word forms store information inside floating text boxes
# (<w:txbxContent>). These are visually positioned on the page and often
# contain user-editable text.
#
# The current assessment forms do NOT use text boxes.
# =============================================================================

    print( "\n=== TEXT BOXES ===" )

    text_boxes = doc.element.xpath( ".//w:txbxContent" )
    print ( f"Text boxes       : {len ( text_boxes ) }" )

    for box_index, box in enumerate( text_boxes ):
        print ( f"\nText box {box_index}" )

        for paragraph in box.xpath(".//w:p" ):
            text = paragraph_text ( paragraph )

            if text.strip():
                print ( f"    {text!r}" )

# =============================================================================
# XML Paragraph Sweep
#
# Walk every XML paragraph (<w:p>) in the document.
#
# This is a lower-level view than doc.paragraphs. It reveals paragraphs that
# python-docx does not necessarily expose as high-level Paragraph objects.
#
# Many reverse-engineering tasks begin by sweeping every paragraph and
# examining the recovered text.
# =============================================================================

    print ( "\n=== CONTENT CONTROLS ===" )

    controls = doc.element.xpath( ".//w:sdt" )
    print ( f"Content controls : {len ( controls ) }" )

# =============================================================================
# Text Boxes
#
# Many Word forms store information inside floating text boxes
# (<w:txbxContent>). These are visually positioned on the page and often
# contain user-editable text.
#
# The current assessment forms do NOT use text boxes.
# =============================================================================

    for control_index, control in enumerate ( controls ):
        text = paragraph_text(control).strip()

# -------------------------------------------------------------------------
# Content-control metadata
#
# Every content control consists of:
#
#     <w:sdt>
#         <w:sdtPr>      properties describing the field
#         <w:sdtContent> user-visible contents
#
# The metadata is optional.
#
# Tag
# ----
# A machine-readable identifier intended for programs.
#
# Example:
#
#     DateField
#     GovernmentParticipants
#     ConfidenceRating
#
# Alias
# -----
# A human-readable label displayed by Word.
#
# Example:
#
#     "Government Participants"
#     "Date"
#     "Overall Confidence"
#
# Some document generators populate these fields, while others leave them
# empty. When present, they provide a convenient semantic description of the
# field being harvested.
# -------------------------------------------------------------------------

        tags =    control.xpath ( "./w:sdtPr/w:tag/@w:val",   namespaces = Word_NameSpace )
        aliases = control.xpath ( "./w:sdtPr/w:alias/@w:val", namespaces = Word_NameSpace )

        if text:
            print ( f"\nControl {control_index}" )
            print ( f"    tag   = {tags}" )
            print ( f"    alias = {aliases}" )
            print ( f"    text  = {text!r}" )


# === PROVENANCE PRINTER ===
def print_provenance():
    current_time = datetime.datetime.now().strftime( "%Y-%m-%d %H:%M:%S" )
    user = os.getlogin()
    notebook_path = Path.cwd()

    print ( "\n=== BASIC PROVENANCE ===" )
    print ( f"Timestamp: {current_time}" )
    print ( f"User: {user}" )
    print ( f"Notebook location: {notebook_path}" )

    print ( "\n=== SYSTEM ===" )
    print ( f"OS: {platform.system()} {platform.release()}" )
    print ( f"Machine: {platform.machine()}" )
    print ( f"Processor: {platform.processor()}" )

    print ( "\n=== PYTHON ===" )
    print ( f"• Version: {platform.python_version()}" )
    print ( f"• Executable: {sys.executable}" )
    print ( f"• Implementation: {platform.python_implementation()}" )


if __name__ == "__main__":
    sweep_document( "bravo.docx" )
    print_provenance( )

# dantopa@quaxolotl:word $ python sweeper-docx.py 

# === HIGH-LEVEL PYTHON-DOCX OBJECTS ===
# Paragraphs       : 70
# Tables           : 0
# Sections         : 1

# === ALL XML PARAGRAPHS ===
#    1: 'Appendix C - IIR #2 Assessment Worksheet'
#    3: 'Reviewer Instructions'
#    4: "This worksheet is intended to document the Government's assessment results for product elements. Before completing this worksheet, reviewers should read Section 9 – Completing the IIR #2 IPT Assessment Worksheet, which provides detailed instructions for completing each section, applying the Confidence Rating Criteria, and documenting evidence-based observations and conclusions."
#    6: 'Date: '
#    7: '3/19/2026'
#    8: 'Control Account: '
#    9: ' Alpha - ABCD'
#   10: 'Product Element Owner Designation: '
#   11: 'Training Systems Branch - NIXZ'
#   12: 'Government Participants: '
#   13: 'Bill Brown (ML)Gary Green (Lead Engineer)Peter Purple (Safe Cracker)Rick Red (Wheel Man)Susan Wood (Look Out)'
#   14: 'Contractor Participants: '
#   15: 'Quint Quartz (CAM)Frank Feldspar (Fence)Oscar Ore (Muscle)Harvey Halide (Muscle)Mike Magnetite (Comic Relief)'
#   17: '1. Leadership Summary & Confidence Ratings'
#   18: 'Top Three Messages for Leadership'
#   19: 'If leadership only remembers three things from this assessment, what should they be?'
#   20: ' 1. Planning Framework is Well EstablishedThe program has developed a coherent planning framework with clearly defined work scope, logical scheduling practices, and consistent alignment across participating organizations. This foundation provides a solid basis for continued planning and execution activities.2. Prepared for the Next Phase of PlanningCurrent planning products indicate sufficient maturity to support the transition into more detailed schedule development and baseline refinement. Remaining issues appear limited in scope and can be addressed during normal planning evolution.3. Maintain Planning Discipline During ExecutionAs planning moves toward execution, continued emphasis on interface management, dependency tracking, and governance processes will help preserve schedule integrity and support effective cross-functional coordination.'
#   22: 'Confidence Ratings'
#   23: 'Overall Confidence Rating (Select One)  '
#   24: 'Strong: ☒Moderate: ☐Low: ☐'
#   26: "The Overall Confidence Rating represents the Government's professional judgment of the contractor's current planning maturity and readiness based on the evidence reviewed during IIR #2. It reflects the Government's point-in-time assessment of the contractor's planning architecture and its readiness to continue progressing toward Integrated Master Schedule (IMS) loading, Performance Measurement Baseline (PMB) implementation, and subsequent IIR activities."
#   28: "Importantly, this rating is not a determination that the IMS or PMB is fully mature, executable, or ready for approval. Instead, it reflects the Government's current confidence that the contractor has established a sufficiently mature planning foundation to support continued planning, integration, and progressive maturation toward Milestone B."
#   30: 'Overall Confidence Rating Rationale'
#   31: " The available evidence supports a Strong confidence assessment. Planning processes, supporting documentation, and management practices demonstrate an appropriate level of maturity for the current phase of the program. Scope definition, schedule development, and product integration activities are generally well organized and exhibit good consistency with program objectives and contractual expectations.The review identified effective coordination across planning disciplines, with interfaces and dependencies documented at a level that supports integrated execution planning. Assessment artifacts were found to be substantially complete and mutually consistent, providing confidence in the overall planning approach. While the team identified several recommendations for continued improvement, these items are incremental in nature and do not materially reduce confidence in the program's planning readiness. Overall, the assessment indicates that the planning foundation is capable of supporting future schedule integration, baseline development, and execution activities."
#   33: 'Next, record the confidence rating for each assessment area using the applicable Confidence Rating Criteria from the IIR #2 Assessment Guide. For each assessment area, provide a brief rationale explaining the basis for the assigned rating. The rationale should summarize the key evidence, observations, strengths, concerns, and uncertainties that most influenced the Government’s judgment. The purpose is not to document every observation, but to clearly explain why the Government reviewers concluded that the assessment area warranted a Strong, Moderate, or Low confidence rating. The rationale should focus on the most significant factors influencing the rating rather than attempting to summarize every discussion, planning artifact, or piece of evidence reviewed.'
#   35: 'Assessment Area 1 – Scope and Deliverables Definition                             Confidence Rating (Select One)  '
#   36: 'Strong: ☒Moderate: ☐Low: ☐'
#   37: ' Review of the submitted planning package indicates a mature definition of the remaining contractual effort. The work breakdown framework and supporting planning documentation establish clear scope boundaries, map deliverables to program-level objectives, and identify organizational ownership for execution. Collectively, the evidence supports a high level of confidence that the contractor has established a disciplined planning foundation and is prepared to proceed with the remaining work.'
#   39: 'Assessment Area 2 – Schedule Planning and IMS-Loading Readiness      Confidence Rating (Select One)  '
#   40: 'Strong: ☐Moderate: ☒Low: ☐'
#   41: ' Examination of the planning package found that the contractor has implemented a structured planning approach capable of supporting subsequent IMS development. Fundamental planning components are present and reasonably mature, demonstrating that the overall architecture has progressed beyond the initial stages. Remaining weaknesses are concentrated in the documentation of planning assumptions, the justification of activity durations, and the definition of cross-functional dependencies. While the overall planning baseline is credible, additional refinement and corroborating evidence are needed before the Government could conclude that the planning process has achieved a fully mature state.'
#   43: 'Assessment Area 3 – Cross-Product/Cross-Function Integration                                            Confidence Rating (Select One)  '
#   44: 'Strong: ☐Moderate: ☐Low: ☒'
#   45: " Evaluation of the submitted planning artifacts identified significant weaknesses in the contractor's cross-product integration strategy, resulting in an overall assessment of Low. Although the planning package reflects preliminary integration activities, it provides insufficient evidence that technical interfaces, schedule dependencies, coordination mechanisms, and organizational responsibilities have been fully identified and incorporated into the planning baseline. The relationship between the assigned product element and the broader program architecture remains only partially defined, limiting confidence that integration can be effectively managed during IMS development. Additional planning detail, interface definition, and cross-organizational coordination will be required before the Government can characterize the contractor's integration planning as mature or execution-ready."
#   47: '2. Products & Artifacts Reviewed'
#   48: 'Document the planning artifacts, technical products, and supporting documentation reviewed by the IPT to support its assessment and confidence determinations.'
#   50: 'The list should identify the primary sources of evidence used during the assessment and demonstrate that conclusions were based on documented information, supplemented by discussions with contractor counterparts as appropriate. When available, include the applicable revision, version, or date of each artifact reviewed.'
#   52: ' 1. Planning Architecture, Version 1.3 (20 May 2026): Used to evaluate the organization of the planning framework, workflow design, and task hierarchy.2. CWBS Dictionary, Revision 4: Examined to validate scope allocation and the definition of work packages.3. Integrated Master Plan (IMP), Revision B (15 May 2026): Reviewed to assess program events, accomplishment criteria, and the overall planning strategy.'
#   54: '3. Strengths Observed'
#   55: "Document the most significant strengths identified during the assessment that increase confidence in the contractor's planning maturity and readiness. Strengths should be supported by objective evidence and highlight effective planning practices, planning artifacts, integration activities, or management disciplines that demonstrate a well-developed planning architecture and increase confidence in the contractor's planning maturity and readiness to support future Integrated Master Schedule (IMS) development, Performance Measurement Baseline (PMB) implementation, and program execution."
#   57: " Strength 1: S–3The planning baseline reflects a comprehensive definition of the remaining contractual scope, organized within a structured product decomposition. The resulting planning architecture establishes a reliable basis for detailed scheduling and future IMS construction.Strength 2: P–5Assessment activities benefited from planning documentation that was complete, organized, and readily traceable to contractor responses. The availability of objective supporting evidence demonstrated a disciplined planning environment and substantiated the Government's overall conclusions."
#   59: '4. Concerns and Risks Observed'
#   60: "Document the most significant concerns identified during the assessment and any associated risks that could reduce confidence in the contractor's planning maturity and readiness or adversely affect future planning, integration, or program execution if not addressed."
#   62: 'Concerns should be supported by objective evidence and identify observed weaknesses, gaps, inconsistencies, or uncertainties that warrant additional management attention. Risks should describe the potential consequences if those concerns remain unresolved. Not every concern creates a program risk, and not every identified risk is currently affecting program execution. The purpose of this section is to identify issues that warrant Government visibility, follow-up, or corrective action as planning continues to mature.'
#   64: 'Definitions'
#   65: "Concern: An observed weakness, gap, inconsistency, or uncertainty that may reduce confidence in the contractor's planning maturity and readiness."
#   66: 'Risk: A potential future consequence that could adversely affect planning, cost, schedule, technical performance, integration, PMB implementation, or program execution if a concern is not addressed.'
#   68: ' Concern 1: S–3The planning baseline demonstrates substantial progress in defining the remaining contractual work, but decomposition of certain planning packages remains incomplete for IMS integration.Risk 1: S–3Failure to complete this decomposition before schedule development may lead to repeated schedule restructuring, weakening confidence in the stability of the execution baseline and delaying PMB implementation.Concern 2: P–3The review identified instances where activity duration estimates were not adequately supported by documented assumptions, historical evidence, or quantitative estimating methods.Risk 2: P–3 If duration estimates are not sufficiently substantiated, the resulting schedule may understate or overstate the effort required, reducing the credibility of execution plans and schedule projections.'
#   70: '5. Open Questions and Follow-Up Items'
#   71: 'Document questions that remain unanswered during the assessment and identify the additional evidence, clarification, documentation, or demonstrations needed from the contractor to complete the evaluation or increase confidence in the assessment findings.'
#   73: "Open Questions and Follow-Up Items identify areas where the available evidence was incomplete, unclear, or inconsistent. These items represent information gaps that should be resolved through additional discussions, demonstrations, or artifact reviews before the Government can fully assess the contractor's planning maturity and readiness. They are not, by themselves, concerns, deficiencies, or risks."
#   75: ' Question 1: S–3Evidence reviewed indicates that external product dependencies have been recognized, but the review team could not fully evaluate the methods used to control and integrate those dependencies throughout schedule development.Follow-Up 1:Provide an end-to-end demonstration of dependency management, including identification, ownership, maintenance, and implementation within the IMS.Question 2: P–3The planning baseline suggests a relationship between the IMP, supporting planning products, and the future schedule; however, that relationship has not yet been validated through objective traceability.Follow-Up 2:Demonstrate how accomplishments defined in the IMP are systematically transformed into IMS activities, logical sequencing, and measurable program milestones.'
#   77: '6. Cross-Product/Cross-Function Integration Observations'
#   78: 'Document observations regarding the extent to which the assessed product element is integrated with other product elements and Government assessment areas, including the identification, coordination, and management of technical, planning, schedule, and deliverable interfaces necessary to support planning maturity and program execution.'
#   80: "Cross-Product/Cross-Function Integration Observations should summarize how effectively the contractor has identified, coordinated, and incorporated cross-product/cross-function interfaces, dependencies, and integration activities into its planning architecture. Observations should be supported by objective evidence and reflect the maturity of cross-functional planning, the visibility of interdependencies, and the Government's confidence that the assigned product element can be successfully integrated into the overall program."
#   82: ' Observation 1: I–5The assessment found that the planning architecture adequately represents interfaces with the Guidance, Avionics, and Ground Systems product elements. Major technical exchanges, deliverable transitions, and coordination responsibilities were consistently described during discussions, indicating a mature understanding of cross-product integration requirements.Observation 2: I–3Although the planning baseline identifies logical relationships with external product elements, the supporting rationale for the timing of those dependencies has not been fully developed. Additional documentation is needed to demonstrate the basis for predecessor and successor sequencing during schedule development.'

# === TEXT BOXES ===
# Text boxes       : 0

# === CONTENT CONTROLS ===
# Content controls : 27

# Control 0
#     tag   = ['Date']
#     alias = ['Date']
#     text  = '3/19/2026'

# Control 1
#     tag   = []
#     alias = []
#     text  = 'Alpha - ABCD'

# Control 2
#     tag   = ['Product']
#     alias = ['Product']
#     text  = 'Training Systems Branch - NIXZ'

# Control 3
#     tag   = []
#     alias = []
#     text  = 'Bill Brown (ML)Gary Green (Lead Engineer)Peter Purple (Safe Cracker)Rick Red (Wheel Man)Susan Wood (Look Out)'

# Control 4
#     tag   = []
#     alias = []
#     text  = 'Quint Quartz (CAM)Frank Feldspar (Fence)Oscar Ore (Muscle)Harvey Halide (Muscle)Mike Magnetite (Comic Relief)'

# Control 5
#     tag   = []
#     alias = []
#     text  = '1. Planning Framework is Well EstablishedThe program has developed a coherent planning framework with clearly defined work scope, logical scheduling practices, and consistent alignment across participating organizations. This foundation provides a solid basis for continued planning and execution activities.2. Prepared for the Next Phase of PlanningCurrent planning products indicate sufficient maturity to support the transition into more detailed schedule development and baseline refinement. Remaining issues appear limited in scope and can be addressed during normal planning evolution.3. Maintain Planning Discipline During ExecutionAs planning moves toward execution, continued emphasis on interface management, dependency tracking, and governance processes will help preserve schedule integrity and support effective cross-functional coordination.'

# Control 6
#     tag   = []
#     alias = []
#     text  = '☒'

# Control 7
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 8
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 9
#     tag   = []
#     alias = []
#     text  = "The available evidence supports a Strong confidence assessment. Planning processes, supporting documentation, and management practices demonstrate an appropriate level of maturity for the current phase of the program. Scope definition, schedule development, and product integration activities are generally well organized and exhibit good consistency with program objectives and contractual expectations.The review identified effective coordination across planning disciplines, with interfaces and dependencies documented at a level that supports integrated execution planning. Assessment artifacts were found to be substantially complete and mutually consistent, providing confidence in the overall planning approach. While the team identified several recommendations for continued improvement, these items are incremental in nature and do not materially reduce confidence in the program's planning readiness. Overall, the assessment indicates that the planning foundation is capable of supporting future schedule integration, baseline development, and execution activities."

# Control 10
#     tag   = []
#     alias = []
#     text  = '☒'

# Control 11
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 12
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 13
#     tag   = []
#     alias = []
#     text  = 'Review of the submitted planning package indicates a mature definition of the remaining contractual effort. The work breakdown framework and supporting planning documentation establish clear scope boundaries, map deliverables to program-level objectives, and identify organizational ownership for execution. Collectively, the evidence supports a high level of confidence that the contractor has established a disciplined planning foundation and is prepared to proceed with the remaining work.'

# Control 14
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 15
#     tag   = []
#     alias = []
#     text  = '☒'

# Control 16
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 17
#     tag   = []
#     alias = []
#     text  = 'Examination of the planning package found that the contractor has implemented a structured planning approach capable of supporting subsequent IMS development. Fundamental planning components are present and reasonably mature, demonstrating that the overall architecture has progressed beyond the initial stages. Remaining weaknesses are concentrated in the documentation of planning assumptions, the justification of activity durations, and the definition of cross-functional dependencies. While the overall planning baseline is credible, additional refinement and corroborating evidence are needed before the Government could conclude that the planning process has achieved a fully mature state.'

# Control 18
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 19
#     tag   = []
#     alias = []
#     text  = '☐'

# Control 20
#     tag   = []
#     alias = []
#     text  = '☒'

# Control 21
#     tag   = []
#     alias = []
#     text  = "Evaluation of the submitted planning artifacts identified significant weaknesses in the contractor's cross-product integration strategy, resulting in an overall assessment of Low. Although the planning package reflects preliminary integration activities, it provides insufficient evidence that technical interfaces, schedule dependencies, coordination mechanisms, and organizational responsibilities have been fully identified and incorporated into the planning baseline. The relationship between the assigned product element and the broader program architecture remains only partially defined, limiting confidence that integration can be effectively managed during IMS development. Additional planning detail, interface definition, and cross-organizational coordination will be required before the Government can characterize the contractor's integration planning as mature or execution-ready."

# Control 22
#     tag   = []
#     alias = []
#     text  = '1. Planning Architecture, Version 1.3 (20 May 2026): Used to evaluate the organization of the planning framework, workflow design, and task hierarchy.2. CWBS Dictionary, Revision 4: Examined to validate scope allocation and the definition of work packages.3. Integrated Master Plan (IMP), Revision B (15 May 2026): Reviewed to assess program events, accomplishment criteria, and the overall planning strategy.'

# Control 23
#     tag   = []
#     alias = []
#     text  = "Strength 1: S–3The planning baseline reflects a comprehensive definition of the remaining contractual scope, organized within a structured product decomposition. The resulting planning architecture establishes a reliable basis for detailed scheduling and future IMS construction.Strength 2: P–5Assessment activities benefited from planning documentation that was complete, organized, and readily traceable to contractor responses. The availability of objective supporting evidence demonstrated a disciplined planning environment and substantiated the Government's overall conclusions."

# Control 24
#     tag   = []
#     alias = []
#     text  = 'Concern 1: S–3The planning baseline demonstrates substantial progress in defining the remaining contractual work, but decomposition of certain planning packages remains incomplete for IMS integration.Risk 1: S–3Failure to complete this decomposition before schedule development may lead to repeated schedule restructuring, weakening confidence in the stability of the execution baseline and delaying PMB implementation.Concern 2: P–3The review identified instances where activity duration estimates were not adequately supported by documented assumptions, historical evidence, or quantitative estimating methods.Risk 2: P–3 If duration estimates are not sufficiently substantiated, the resulting schedule may understate or overstate the effort required, reducing the credibility of execution plans and schedule projections.'

# Control 25
#     tag   = []
#     alias = []
#     text  = 'Question 1: S–3Evidence reviewed indicates that external product dependencies have been recognized, but the review team could not fully evaluate the methods used to control and integrate those dependencies throughout schedule development.Follow-Up 1:Provide an end-to-end demonstration of dependency management, including identification, ownership, maintenance, and implementation within the IMS.Question 2: P–3The planning baseline suggests a relationship between the IMP, supporting planning products, and the future schedule; however, that relationship has not yet been validated through objective traceability.Follow-Up 2:Demonstrate how accomplishments defined in the IMP are systematically transformed into IMS activities, logical sequencing, and measurable program milestones.'

# Control 26
#     tag   = []
#     alias = []
#     text  = 'Observation 1: I–5The assessment found that the planning architecture adequately represents interfaces with the Guidance, Avionics, and Ground Systems product elements. Major technical exchanges, deliverable transitions, and coordination responsibilities were consistently described during discussions, indicating a mature understanding of cross-product integration requirements.Observation 2: I–3Although the planning baseline identifies logical relationships with external product elements, the supporting rationale for the timing of those dependencies has not been fully developed. Additional documentation is needed to demonstrate the basis for predecessor and successor sequencing during schedule development.'

# === BASIC PROVENANCE ===
# Timestamp: 2026-08-06 19:56:56
# User: dantopa
# Notebook location: /Users/dantopa/repos-quaxolotl/github/jop/python/bae/word

# === SYSTEM ===
# OS: Darwin 25.5.0
# Machine: x86_64
# Processor: i386

# === PYTHON ===
# • Version: 3.13.14
# • Executable: /opt/local/bin/python
# • Implementation: CPython