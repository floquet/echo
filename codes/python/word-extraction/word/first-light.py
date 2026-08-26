#! /opt/local/bin/python3.13

import os
import sys
import platform
import datetime
from pathlib import Path

from docx import Document

# === PROVENANCE PRINTER ===
def print_provenance():
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    user = os.getlogin()
    notebook_path = Path.cwd()

    print("\n=== BASIC PROVENANCE ===")
    print(f"Timestamp: {current_time}")
    print(f"User: {user}")
    print(f"Notebook location: {notebook_path}")

    print("\n=== SYSTEM ===")
    print(f"OS: {platform.system()} {platform.release()}")
    print(f"Machine: {platform.machine()}")
    print(f"Processor: {platform.processor()}")

    print("\n=== PYTHON ===")
    print(f"• Version: {platform.python_version()}")
    print(f"• Executable: {sys.executable}")
    print(f"• Implementation: {platform.python_implementation()}")

#  ==   ==   == ==   ==   == ==   ==   == ==   ==   ==  #

if __name__ == "__main__":


    doc = Document()

    doc.add_heading ('Greeting', level = 1 )
    doc.add_paragraph ("Would you like to play a game, Dave?" )

    output_file = "wargames.docx"
    doc.save ( output_file ) 

    print( f"Wrote {output_file}" )

    print_provenance()

# dantopa@quaxolotl:word $ python first-light.py 
# Wrote wargames.docx

# === BASIC PROVENANCE ===
# Timestamp: 2026-08-06 15:44:11
# User: dantopa
# Notebook location: /Users/dantopa/repos-quaxolotl/github/jop/python/bae/word

# === SYSTEM ===
# OS: Darwin 25.5.0
# Machine: x86_64
# Processor: i386

# === PYTHON ===
# • Version: 3.13.14
# • Executable: /opt/local/bin/python
# • Implementation: CPython
