#! /opt/local/bin/python3.13

# Daniel Topa
# BAE 899149

import os
import sys
import platform
import datetime
from pathlib import Path

from docx import Document

# === PROVENANCE PRINTER ===
def print_provenance():
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S" )
    user = os.getlogin()
    notebook_path = Path.cwd()

    print ( "\n=== BASIC PROVENANCE ===" )
    print ( f"Timestamp: {current_time}" )
    print ( f"User: {user}" )
    print ( f"Notebook location: {notebook_path}" )

    print ( "\n=== SYSTEM ===" )
    print ( f"OS: {platform.system()} {platform.release()}" )
    print ( f"Machine: {platform.machine()}" )
    print ( f"Processor: {platform.processor()}" )

    print ( "\n=== PYTHON ===" )
    print ( f"• Version: {platform.python_version()}" )
    print ( f"• Executable: {sys.executable}" )
    print ( f"• Implementation: {platform.python_implementation()}" )

#  ==   ==   == ==   ==   == ==   ==   == ==   ==   ==  #

if __name__ == "__main__":

    doc = Document()

    input_file = "bravo.docx"
    doc = Document( input_file )
    print( f"Ingested {input_file}" )

    # *.docx queries
    print ( f"Document queries: " )
    print ( f"Number of tables: ", len( doc.tables ) )
    print ( f"Niumber of sections: ", len( doc.sections ) )
    print ( )
    print ( f"Paragraph queries start: " )
    print ( f"Paragraph lengths, characters: { len ( doc.paragraphs ) }" )
    # i: counter
    # p: text
    print ( f"Sweep paragraph properties" )
    for i, p in enumerate ( doc.paragraphs ) :
        print ( f"{i:4d}: { p.text!r }" )
        # print(p.text)
        # print ( f"style name = ", p.style.name )
        # print ( f"alignment  = ", p.alignment )

    print ( ) 
    print ( f"Looking deeper for user input fields")
    for i, node in enumerate( doc.element.xpath(".//w:t") ):
        print ( f"{i:4d}: { node.text!r }" )

    print ( ) 
    print ( f"Stitch content together")
    paragraph_elements = doc.element.xpath(".//w:p")

    for i, paragraph in enumerate( paragraph_elements ):
        fragments = paragraph.xpath(".//w:t/text()")
        text = "".join( fragments )

    if text.strip():
        print ( f"{i:4d}: {text!r}" )

    print_provenance()

# dantopa@quaxolotl:word $ python alpha-reader.py 
# Ingested bravo.docx
# Document queries: 
# Number of tables:  0
# Niumber of sections:  1

# Paragraph queries start: 
# Paragraph lengths, characters: 70
# Sweep paragraph properties
#    0: ''
#    1: 'Appendix C - IIR #2 Assessment Worksheet'
#    2: ''
#    3: 'Reviewer Instructions'
#    4: "This worksheet is intended to document the Government's assessment results for product elements. Before completing this worksheet, reviewers should read Section 9 – Completing the IIR #2 IPT Assessment Worksheet, which provides detailed instructions for completing each section, applying the Confidence Rating Criteria, and documenting evidence-based observations and conclusions."
#    5: ''
#    6: 'Date: '
#    7: 'Control Account: '
#    8: 'Product Element Owner Designation: '
#    9: 'Government Participants: '
#   10: 'Contractor Participants: '
#   11: ''
#   12: '1. Leadership Summary & Confidence Ratings'
#   13: 'Top Three Messages for Leadership'
#   14: 'If leadership only remembers three things from this assessment, what should they be?'
#   15: ''
#   16: 'Confidence Ratings'
#   17: 'Overall Confidence Rating (Select One)  '
#   18: 'Strong: \tModerate: \tLow: '
#   19: ''
#   20: "The Overall Confidence Rating represents the Government's professional judgment of the contractor's current planning maturity and readiness based on the evidence reviewed during IIR #2. It reflects the Government's point-in-time assessment of the contractor's planning architecture and its readiness to continue progressing toward Integrated Master Schedule (IMS) loading, Performance Measurement Baseline (PMB) implementation, and subsequent IIR activities."
#   21: ''
#   22: "Importantly, this rating is not a determination that the IMS or PMB is fully mature, executable, or ready for approval. Instead, it reflects the Government's current confidence that the contractor has established a sufficiently mature planning foundation to support continued planning, integration, and progressive maturation toward Milestone B."
#   23: ''
#   24: 'Overall Confidence Rating Rationale'
#   25: ''
#   26: 'Next, record the confidence rating for each assessment area using the applicable Confidence Rating Criteria from the IIR #2 Assessment Guide. For each assessment area, provide a brief rationale explaining the basis for the assigned rating. The rationale should summarize the key evidence, observations, strengths, concerns, and uncertainties that most influenced the Government’s judgment. The purpose is not to document every observation, but to clearly explain why the Government reviewers concluded that the assessment area warranted a Strong, Moderate, or Low confidence rating. The rationale should focus on the most significant factors influencing the rating rather than attempting to summarize every discussion, planning artifact, or piece of evidence reviewed.'
#   27: ''
#   28: 'Assessment Area 1 – Scope and Deliverables Definition                             Confidence Rating (Select One)  '
#   29: 'Strong: \tModerate: \tLow: '
#   30: ''
#   31: 'Assessment Area 2 – Schedule Planning and IMS-Loading Readiness      Confidence Rating (Select One)  '
#   32: 'Strong: \tModerate: \tLow: '
#   33: ''
#   34: 'Assessment Area 3 – Cross-Product/Cross-Function Integration                                            Confidence Rating (Select One)  '
#   35: 'Strong: \tModerate: \tLow: '
#   36: ''
#   37: '2. Products & Artifacts Reviewed'
#   38: 'Document the planning artifacts, technical products, and supporting documentation reviewed by the IPT to support its assessment and confidence determinations.'
#   39: ''
#   40: 'The list should identify the primary sources of evidence used during the assessment and demonstrate that conclusions were based on documented information, supplemented by discussions with contractor counterparts as appropriate. When available, include the applicable revision, version, or date of each artifact reviewed.'
#   41: ''
#   42: ''
#   43: '3. Strengths Observed'
#   44: "Document the most significant strengths identified during the assessment that increase confidence in the contractor's planning maturity and readiness. Strengths should be supported by objective evidence and highlight effective planning practices, planning artifacts, integration activities, or management disciplines that demonstrate a well-developed planning architecture and increase confidence in the contractor's planning maturity and readiness to support future Integrated Master Schedule (IMS) development, Performance Measurement Baseline (PMB) implementation, and program execution."
#   45: ''
#   46: ''
#   47: '4. Concerns and Risks Observed'
#   48: "Document the most significant concerns identified during the assessment and any associated risks that could reduce confidence in the contractor's planning maturity and readiness or adversely affect future planning, integration, or program execution if not addressed."
#   49: ''
#   50: 'Concerns should be supported by objective evidence and identify observed weaknesses, gaps, inconsistencies, or uncertainties that warrant additional management attention. Risks should describe the potential consequences if those concerns remain unresolved. Not every concern creates a program risk, and not every identified risk is currently affecting program execution. The purpose of this section is to identify issues that warrant Government visibility, follow-up, or corrective action as planning continues to mature.'
#   51: ''
#   52: 'Definitions'
#   53: "Concern: An observed weakness, gap, inconsistency, or uncertainty that may reduce confidence in the contractor's planning maturity and readiness."
#   54: 'Risk: A potential future consequence that could adversely affect planning, cost, schedule, technical performance, integration, PMB implementation, or program execution if a concern is not addressed.'
#   55: ''
#   56: ''
#   57: '5. Open Questions and Follow-Up Items'
#   58: 'Document questions that remain unanswered during the assessment and identify the additional evidence, clarification, documentation, or demonstrations needed from the contractor to complete the evaluation or increase confidence in the assessment findings.'
#   59: ''
#   60: "Open Questions and Follow-Up Items identify areas where the available evidence was incomplete, unclear, or inconsistent. These items represent information gaps that should be resolved through additional discussions, demonstrations, or artifact reviews before the Government can fully assess the contractor's planning maturity and readiness. They are not, by themselves, concerns, deficiencies, or risks."
#   61: ''
#   62: ''
#   63: '6. Cross-Product/Cross-Function Integration Observations'
#   64: 'Document observations regarding the extent to which the assessed product element is integrated with other product elements and Government assessment areas, including the identification, coordination, and management of technical, planning, schedule, and deliverable interfaces necessary to support planning maturity and program execution.'
#   65: ''
#   66: "Cross-Product/Cross-Function Integration Observations should summarize how effectively the contractor has identified, coordinated, and incorporated cross-product/cross-function interfaces, dependencies, and integration activities into its planning architecture. Observations should be supported by objective evidence and reflect the maturity of cross-functional planning, the visibility of interdependencies, and the Government's confidence that the assigned product element can be successfully integrated into the overall program."
#   67: ''
#   68: ''
#   69: ''

# Looking deeper for user input fields
#    0: 'Appendix C - IIR #2 Assessment Worksheet'
#    1: 'Reviewer Instructions'
#    2: "This worksheet is intended to document the Government's assessment results for product element"
#    3: 's'
#    4: '. Before completing this worksheet, reviewers should read '
#    5: 'Section 9 – Completing the IIR #2 IPT Assessment Worksheet'
#    6: ', which provides detailed instructions for completing each section, applying the Confidence Rating Criteria, and documenting evidence-based observations and conclusions.'
#    7: 'Date:'
#    8: ' '
#    9: '3/19/2026'
#   10: 'Control Account'
#   11: ':'
#   12: ' '
#   13: ' '
#   14: 'Alpha - ABCD'
#   15: 'Product Elemen'
#   16: 't Owner Designation'
#   17: ':'
#   18: ' '
#   19: 'Training Systems Branch - NIXZ'
#   20: 'Government Participants:'
#   21: ' '
#   22: 'Bill Brown'
#   23: ' (ML)'
#   24: 'Gary Green'
#   25: ' (Lead Engineer)'
#   26: 'Peter Purple'
#   27: ' ('
#   28: 'Safe Cracker'
#   29: ')'
#   30: 'Rick Red'
#   31: ' ('
#   32: 'Wheel Man'
#   33: ')'
#   34: 'Susan Wood ('
#   35: 'Look Out'
#   36: ')'
#   37: 'Contractor Participants:'
#   38: ' '
#   39: 'Quint Quartz'
#   40: ' (CAM)'
#   41: 'Frank Feldspar'
#   42: ' ('
#   43: 'Fence'
#   44: ')'
#   45: 'Oscar Ore'
#   46: ' ('
#   47: 'Muscle'
#   48: ')'
#   49: 'Harvey Halide'
#   50: ' ('
#   51: 'Muscle'
#   52: ')'
#   53: 'Mike Magnetite'
#   54: ' ('
#   55: 'Comic Relief'
#   56: ')'
#   57: '1. Leadership Summary & Confidence Ratings'
#   58: 'Top Three Messages for Leadership'
#   59: 'If leadership only remembers three things from this assessment, what should they be?'
#   60: ' 1. Planning Framework is Well Established'
#   61: 'The program has developed a coherent planning framework with clearly defined work scope, logical scheduling practices, and consistent alignment across participating organizations. This foundation provides a solid basis for continued planning and execution activities.'
#   62: '2. Prepared for the Next Phase of Planning'
#   63: 'Current planning products indicate sufficient maturity to support the transition into more detailed schedule development and baseline refinement. Remaining issues appear limited in scope and can be addressed during normal planning evolution.'
#   64: '3. Maintain Planning Discipline During Execution'
#   65: 'As planning moves toward execution, continued emphasis on interface management, dependency tracking, and governance processes will help preserve schedule integrity and support effective cross-functional coordination.'
#   66: 'Confidence Ratings'
#   67: 'Overall Confidence Rating (Select One)  '
#   68: 'Strong:'
#   69: ' '
#   70: '☒'
#   71: 'Moderate:'
#   72: ' '
#   73: '☐'
#   74: 'Low:'
#   75: ' '
#   76: '☐'
#   77: "The Overall Confidence Rating represents the Government's professional judgment of the contractor's current planning maturity and readiness based on the evidence reviewed during IIR #2. It reflects the Government's point-in-time assessment of the contractor's planning architecture and its readiness to continue progressing toward Integrated Master Schedule (IMS) loading, Performance Measurement Baseline (PMB) implementation, and subsequent IIR activities."
#   78: "Importantly, this rating is not a determination that the IMS or PMB is fully mature, executable, or ready for approval. Instead, it reflects the Government's current confidence that the contractor has established a sufficiently mature planning foundation to support continued planning, integration, and progressive maturation toward Milestone B."
#   79: 'Overall Confidence Rating '
#   80: 'Rationale'
#   81: ' The available evidence supports a Strong confidence assessment. Planning processes, supporting documentation, and management practices demonstrate an appropriate level of maturity for the current phase of the program. Scope definition, schedule development, and product integration activities are generally well organized and exhibit good consistency with program objectives and contractual expectations.'
#   82: "The review identified effective coordination across planning disciplines, with interfaces and dependencies documented at a level that supports integrated execution planning. Assessment artifacts were found to be substantially complete and mutually consistent, providing confidence in the overall planning approach. While the team identified several recommendations for continued improvement, these items are incremental in nature and do not materially reduce confidence in the program's planning readiness. Overall, the assessment indicates that the planning foundation is capable of supporting future schedule integration, baseline development, and execution activities."
#   83: 'Next, record'
#   84: ' the confidence rating for each assessment area using the applicable '
#   85: 'Confidence Rating Criteria'
#   86: ' '
#   87: 'from the IIR #2 Assessment Guide. For each assessment area, provide a brief rationale explaining the basis for the assigned rating. The rationale should summarize the key evidence, observations, strengths, concerns, and uncertainties that most influenced the '
#   88: 'Government’s '
#   89: 'judgment. The purpose is not to document every observation, but to clearly explain why the '
#   90: 'Government reviewers'
#   91: ' concluded that the assessment area warranted a '
#   92: 'Strong, Moderate, or Low'
#   93: ' confidence rating. The rationale should focus on the most significant factors influencing the rating rather than attempting to summarize every discussion, planning artifact, or piece of evidence reviewed.'
#   94: 'Assessment Area 1 – Scope and Deliverables Definition'
#   95: ' '
#   96: '                            '
#   97: 'Confidence Rating'
#   98: ' '
#   99: '(Select One)  '
#  100: 'Strong:'
#  101: ' '
#  102: '☒'
#  103: 'Moderate:'
#  104: ' '
#  105: '☐'
#  106: 'Low:'
#  107: ' '
#  108: '☐'
#  109: ' Review of the submitted planning package indicates a mature definition of the remaining contractual effort. The work breakdown framework and supporting planning documentation establish clear scope boundaries, map deliverables to program-level objectives, and identify organizational ownership for execution. Collectively, the evidence supports a high level of confidence that the contractor has established a disciplined planning foundation and is prepared to proceed with the remaining work.'
#  110: 'Assessment Area 2 – Schedule Planning and IMS-Loading Readiness      '
#  111: 'Confidence Rating'
#  112: ' '
#  113: '(Select One)  '
#  114: 'Strong:'
#  115: ' '
#  116: '☐'
#  117: 'Moderate:'
#  118: ' '
#  119: '☒'
#  120: 'Low:'
#  121: ' '
#  122: '☐'
#  123: ' Examination of the planning package found that the contractor has implemented a structured planning approach capable of supporting subsequent IMS development. Fundamental planning components are present and reasonably mature, demonstrating that the overall architecture has progressed beyond the initial stages. Remaining weaknesses are concentrated in the documentation of planning assumptions, the justification of activity durations, and the definition of cross-functional dependencies. While the overall planning baseline is credible, additional refinement and corroborating evidence are needed before the Government could conclude that the planning process has achieved a fully mature state.'
#  124: 'Assessment Area 3 – Cross-Product'
#  125: '/Cross-Function'
#  126: ' Integration                                            '
#  127: 'Confidence Rating'
#  128: ' '
#  129: '(Select One)  '
#  130: 'Strong:'
#  131: ' '
#  132: '☐'
#  133: 'Moderate:'
#  134: ' '
#  135: '☐'
#  136: 'Low:'
#  137: ' '
#  138: '☒'
#  139: " Evaluation of the submitted planning artifacts identified significant weaknesses in the contractor's cross-product integration strategy, resulting in an overall assessment of Low. Although the planning package reflects preliminary integration activities, it provides insufficient evidence that technical interfaces, schedule dependencies, coordination mechanisms, and organizational responsibilities have been fully identified and incorporated into the planning baseline. The relationship between the assigned product element and the broader program architecture remains only partially defined, limiting confidence that integration can be effectively managed during IMS development. Additional planning detail, interface definition, and cross-organizational coordination will be required before the Government can characterize the contractor's integration planning as mature or execution-ready."
#  140: '2'
#  141: '. Products '
#  142: '& Artifacts '
#  143: 'Reviewed'
#  144: 'Document the planning artifacts, technical products, and supporting documentation reviewed by the IPT to support its assessment and confidence determinations.'
#  145: 'The list should identify the primary sources of evidence used during the assessment and demonstrate that conclusions were based on documented information, supplemented by '
#  146: 'discussions with contractor counterparts as appropriate. When available, include the applicable revision, version, or date of each artifact reviewed.'
#  147: ' 1. Planning Architecture, Version 1.3 (20 May 2026): Used to evaluate the organization of the planning framework, workflow design, and task hierarchy.'
#  148: '2. CWBS Dictionary, Revision 4: Examined to validate scope allocation and the definition of work packages.'
#  149: '3. Integrated Master Plan (IMP), Revision B (15 May 2026): Reviewed to assess program events, accomplishment criteria, and the overall planning strategy.'
#  150: '3. Strengths Observed'
#  151: "Document the most significant strengths identified during the assessment that increase confidence in the contractor's planning maturity and readiness. Strengths should be supported by objective evidence and highlight effective planning practices, planning artifacts, integration activities, or management disciplines that demonstrate a well-developed planning architecture and increase confidence in the contractor's planning maturity and readiness to support future Integrated Master Schedule (IMS) development, Performance Measurement Baseline (PMB) implementation, and program execution."
#  152: ' Strength 1: S–3'
#  153: 'The planning baseline reflects a comprehensive definition of the remaining contractual scope, organized within a structured product decomposition. The resulting planning architecture establishes a reliable basis for detailed scheduling and future IMS construction.'
#  154: 'Strength 2: P–5'
#  155: "Assessment activities benefited from planning documentation that was complete, organized, and readily traceable to contractor responses. The availability of objective supporting evidence demonstrated a disciplined planning environment and substantiated the Government's overall conclusions."
#  156: '4. Concerns and Risks Observed'
#  157: "Document the most significant concerns identified during the assessment and any associated risks that could reduce confidence in the contractor's planning maturity and readiness or adversely affect future planning, integration, or program execution if not addressed."
#  158: 'Concerns should be supported by objective evidence and identify observed weaknesses, gaps, inconsistencies, or uncertainties that warrant additional management attention. Risks should describe the potential consequences if those concerns remain unresolved.'
#  159: ' '
#  160: 'Not every concern creates a program risk, and not every identified risk is currently affecting program execution. The purpose of this section is to identify issues that warrant Government visibility, follow-up, or corrective action as planning continues to mature.'
#  161: 'Definitions'
#  162: "Concern: An observed weakness, gap, inconsistency, or uncertainty that may reduce confidence in the contractor's planning maturity and readiness."
#  163: 'Risk: A potential future consequence that could adversely affect planning, cost, schedule, technical performance, integration, PMB implementation, or program execution if a concern is not addressed.'
#  164: ' Concern 1: S–3'
#  165: 'The planning baseline demonstrates substantial progress in defining the remaining contractual work, but decomposition of certain planning packages remains incomplete for IMS integration.'
#  166: 'Risk 1: S–3'
#  167: 'Failure to complete this decomposition before schedule development may lead to repeated schedule restructuring, weakening confidence in the stability of the execution baseline and delaying PMB implementation.'
#  168: 'Concern 2: P–3'
#  169: 'The review identified instances where activity duration estimates were not adequately supported by documented assumptions, historical evidence, or quantitative estimating methods.'
#  170: 'Risk 2: P–3'
#  171: ' '
#  172: 'If duration estimates are not sufficiently substantiated, the resulting schedule may understate or overstate the effort required, reducing the credibility of execution plans and schedule projections.'
#  173: '5. Open Questions and Follow-Up Items'
#  174: 'Document questions that remain unanswered during the assessment and identify the additional evidence, clarification, documentation, or demonstrations needed from the contractor to complete the evaluation or increase confidence in the assessment findings.'
#  175: 'Open Questions and Follow-Up Items identify areas where the available evidence was incomplete, unclear, or inconsistent. These items represent information gaps that should '
#  176: "be resolved through additional discussions, demonstrations, or artifact reviews before the Government can fully assess the contractor's planning maturity and readiness. They are not, by themselves, concerns, deficiencies, or risks."
#  177: ' Question 1: S–3'
#  178: 'Evidence reviewed indicates that external product dependencies have been recognized, but the review team could not fully evaluate the methods used to control and integrate those dependencies throughout schedule development.'
#  179: 'Follow-Up 1:'
#  180: 'Provide an end-to-end demonstration of dependency management, including identification, ownership, maintenance, and implementation within the IMS.'
#  181: 'Question 2: P–3'
#  182: 'The planning baseline suggests a relationship between the IMP, supporting planning products, and the future schedule; however, that relationship has not yet been validated through objective traceability.'
#  183: 'Follow-Up 2:'
#  184: 'Demonstrate how accomplishments defined in the IMP are systematically transformed into IMS activities, logical sequencing, and measurable program milestones.'
#  185: '6. Cross-'
#  186: 'Product'
#  187: '/Cross-Function'
#  188: ' Integration Observations'
#  189: 'Document observations regarding the extent to which the assessed product element is integrated with other product elements and Government assessment areas, including the identification, coordination, and management of technical, planning, schedule, and deliverable interfaces necessary to support planning maturity and program execution.'
#  190: 'Cross-Product'
#  191: '/Cross-Function'
#  192: ' Integration Observations should summarize how effectively the contractor has identified, coordinated, and incorporated cross-product'
#  193: '/cross-function'
#  194: " interfaces, dependencies, and integration activities into its planning architecture. Observations should be supported by objective evidence and reflect the maturity of cross-functional planning, the visibility of interdependencies, and the Government's confidence that the assigned product element can be successfully integrated into the overall program."
#  195: ' Observation 1: I–5'
#  196: 'The assessment found that the planning architecture adequately represents interfaces with the Guidance, Avionics, and Ground Systems product elements. Major technical exchanges, deliverable transitions, and coordination responsibilities were consistently described during discussions, indicating a mature understanding of cross-product integration requirements.'
#  197: 'Observation 2: I–3'
#  198: 'Although the planning baseline identifies logical relationships with external product elements, the supporting rationale for the timing of those dependencies has not been fully developed. Additional documentation is needed to demonstrate the basis for predecessor and successor sequencing during schedule development.'

# Stitch content together

# === BASIC PROVENANCE ===
# Timestamp: 2026-08-06 16:39:01
# User: dantopa
# Notebook location: /Users/dantopa/repos-quaxolotl/github/jop/python/bae/word

# === SYSTEM ===
# OS: Darwin 25.5.0
# Machine: x86_64
# Processor: i386

# === PYTHON ===
# • Version: 3.13.14
# • Executable: /opt/local/bin/python
# • Implementation: CPython